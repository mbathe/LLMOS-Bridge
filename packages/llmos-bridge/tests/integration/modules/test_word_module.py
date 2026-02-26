"""Integration tests — WordModule against real temp .docx files."""

from __future__ import annotations

from pathlib import Path

import pytest
import docx

from llmos_bridge.modules.word import WordModule


@pytest.fixture
def module() -> WordModule:
    return WordModule()


@pytest.fixture
def doc_path(tmp_path: Path) -> Path:
    """Create a simple Word document with some content."""
    doc = docx.Document()
    doc.add_paragraph("Hello World", style="Normal")
    doc.add_paragraph("Second paragraph", style="Normal")
    doc.add_paragraph("Third paragraph", style="Normal")
    path = tmp_path / "test.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def doc_with_table(tmp_path: Path) -> Path:
    """Create a Word document with a table."""
    doc = docx.Document()
    doc.add_paragraph("Document with table")
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(0, 2).text = "C"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"
    table.cell(1, 2).text = "3"
    path = tmp_path / "table.docx"
    doc.save(str(path))
    return path


@pytest.mark.integration
class TestDocumentLifecycle:
    async def test_create_document(self, module: WordModule, tmp_path: Path) -> None:
        out_path = tmp_path / "new.docx"
        result = await module._action_create_document(
            {"output_path": str(out_path), "title": "Test Doc", "author": "Tester"}
        )
        assert result["created"] is True
        assert out_path.exists()

    async def test_open_document(self, module: WordModule, doc_path: Path) -> None:
        result = await module._action_open_document({"path": str(doc_path)})
        assert result["paragraph_count"] >= 3

    async def test_save_document(self, module: WordModule, doc_path: Path) -> None:
        result = await module._action_save_document({"path": str(doc_path)})
        assert result["saved"] is True

    async def test_save_document_to_new_path(self, module: WordModule, doc_path: Path, tmp_path: Path) -> None:
        out_path = tmp_path / "saved_copy.docx"
        result = await module._action_save_document(
            {"path": str(doc_path), "output_path": str(out_path)}
        )
        assert result["saved"] is True
        assert out_path.exists()

    async def test_get_document_meta(self, module: WordModule, doc_path: Path) -> None:
        result = await module._action_get_document_meta({"path": str(doc_path)})
        assert "path" in result
        assert "title" in result
        assert "author" in result

    async def test_set_document_properties(self, module: WordModule, doc_path: Path) -> None:
        result = await module._action_set_document_properties(
            {"path": str(doc_path), "title": "New Title", "author": "New Author"}
        )
        assert result["updated"] is True
        # Verify with fresh module
        module2 = WordModule()
        meta = await module2._action_get_document_meta({"path": str(doc_path)})
        assert meta["title"] == "New Title"


@pytest.mark.integration
class TestReadOperations:
    async def test_read_document(self, module: WordModule, doc_path: Path) -> None:
        result = await module._action_read_document({"path": str(doc_path)})
        assert "paragraphs" in result
        texts = [p["text"] for p in result["paragraphs"]]
        assert "Hello World" in texts

    async def test_read_document_with_tables(self, module: WordModule, doc_with_table: Path) -> None:
        result = await module._action_read_document(
            {"path": str(doc_with_table), "include_tables": True}
        )
        assert "tables" in result
        assert len(result["tables"]) >= 1

    async def test_list_paragraphs(self, module: WordModule, doc_path: Path) -> None:
        result = await module._action_list_paragraphs({"path": str(doc_path)})
        assert result["count"] >= 3
        assert all("text" in p and "style" in p for p in result["paragraphs"])

    async def test_list_paragraphs_exclude_empty(self, module: WordModule, doc_path: Path) -> None:
        result = await module._action_list_paragraphs(
            {"path": str(doc_path), "include_empty": False}
        )
        for para in result["paragraphs"]:
            assert para["text"].strip() != ""

    async def test_list_tables(self, module: WordModule, doc_with_table: Path) -> None:
        result = await module._action_list_tables({"path": str(doc_with_table)})
        assert result["count"] >= 1
        assert result["tables"][0]["rows"] == 2
        assert result["tables"][0]["cols"] == 3

    async def test_extract_text(self, module: WordModule, doc_path: Path) -> None:
        result = await module._action_extract_text({"path": str(doc_path)})
        assert "Hello World" in result["text"]
        assert result["length"] > 0

    async def test_count_words(self, module: WordModule, doc_path: Path) -> None:
        result = await module._action_count_words({"path": str(doc_path)})
        assert result["word_count"] >= 3  # "Hello World" + "Second" + etc


@pytest.mark.integration
class TestParagraphOperations:
    async def test_write_paragraph(self, module: WordModule, doc_path: Path) -> None:
        result = await module._action_write_paragraph(
            {"path": str(doc_path), "text": "New paragraph added"}
        )
        assert "paragraph_index" in result
        # Verify it was saved
        module2 = WordModule()
        content = await module2._action_extract_text({"path": str(doc_path)})
        assert "New paragraph added" in content["text"]

    async def test_write_paragraph_with_style(self, module: WordModule, doc_path: Path) -> None:
        result = await module._action_write_paragraph(
            {"path": str(doc_path), "text": "Heading text", "style": "Heading 1"}
        )
        assert result["style"] == "Heading 1"

    async def test_write_bold_paragraph(self, module: WordModule, doc_path: Path) -> None:
        result = await module._action_write_paragraph(
            {"path": str(doc_path), "text": "Bold text", "bold": True}
        )
        assert "paragraph_index" in result

    async def test_insert_page_break(self, module: WordModule, doc_path: Path) -> None:
        result = await module._action_insert_page_break({"path": str(doc_path)})
        assert result["page_break_inserted"] is True

    async def test_find_replace(self, module: WordModule, doc_path: Path) -> None:
        result = await module._action_find_replace(
            {"path": str(doc_path), "find": "Hello", "replace": "Hi"}
        )
        assert result["replacements_made"] >= 1
        module2 = WordModule()
        content = await module2._action_extract_text({"path": str(doc_path)})
        assert "Hi World" in content["text"]
        assert "Hello World" not in content["text"]

    async def test_delete_paragraph(self, module: WordModule, doc_path: Path) -> None:
        result = await module._action_delete_paragraph(
            {"path": str(doc_path), "paragraph_index": 0}
        )
        assert "deleted_index" in result


@pytest.mark.integration
class TestTableOperations:
    async def test_insert_table(self, module: WordModule, doc_path: Path) -> None:
        result = await module._action_insert_table(
            {
                "path": str(doc_path),
                "rows": 3,
                "cols": 2,
                "data": [["H1", "H2"], ["r1c1", "r1c2"], ["r2c1", "r2c2"]],
            }
        )
        assert result["rows"] == 3
        assert result["cols"] == 2

    async def test_modify_table_cell(self, module: WordModule, doc_with_table: Path) -> None:
        result = await module._action_modify_table_cell(
            {
                "path": str(doc_with_table),
                "table_index": 0,
                "row": 0,
                "col": 0,
                "text": "Modified",
            }
        )
        assert result["updated"] is True

    async def test_add_table_row(self, module: WordModule, doc_with_table: Path) -> None:
        result = await module._action_add_table_row(
            {
                "path": str(doc_with_table),
                "table_index": 0,
                "data": ["X", "Y", "Z"],
            }
        )
        assert "new_row_index" in result

    async def test_insert_list(self, module: WordModule, doc_path: Path) -> None:
        result = await module._action_insert_list(
            {
                "path": str(doc_path),
                "items": ["Item 1", "Item 2", "Item 3"],
                "list_style": "bullet",
            }
        )
        assert result["items_inserted"] == 3


@pytest.mark.integration
class TestFormatting:
    async def test_set_margins(self, module: WordModule, doc_path: Path) -> None:
        result = await module._action_set_margins(
            {
                "path": str(doc_path),
                "top": 2.54,
                "bottom": 2.54,
                "left": 3.17,
                "right": 3.17,
            }
        )
        assert result["top_cm"] == 2.54

    async def test_apply_style(self, module: WordModule, doc_path: Path) -> None:
        result = await module._action_apply_style(
            {"path": str(doc_path), "paragraph_index": 0, "style": "Heading 2"}
        )
        assert result["style"] == "Heading 2"

    async def test_set_default_font(self, module: WordModule, doc_path: Path) -> None:
        result = await module._action_set_default_font(
            {"path": str(doc_path), "font_name": "Arial", "font_size": 11}
        )
        assert result["font_name"] == "Arial"


# ---------------------------------------------------------------------------
# Extended formatting & rich-content tests
# ---------------------------------------------------------------------------


@pytest.mark.integration
class TestFormatText:
    async def test_format_text_bold(self, module: WordModule, doc_path: Path) -> None:
        """Write a paragraph with a run, then format that run bold."""
        # Write a paragraph so we know it has a single run.
        await module._action_write_paragraph(
            {"path": str(doc_path), "text": "FormattableText"}
        )
        # Find the index of the newly added paragraph.
        info = await module._action_list_paragraphs({"path": str(doc_path)})
        texts = [p["text"] for p in info["paragraphs"]]
        idx = texts.index("FormattableText")

        result = await module._action_format_text(
            {
                "path": str(doc_path),
                "paragraph_index": idx,
                "run_index": 0,
                "bold": True,
                "italic": True,
            }
        )
        assert result["formatted"] is True
        assert result["paragraph_index"] == idx
        assert result["run_index"] == 0

    async def test_format_text_font_name_and_size(self, module: WordModule, doc_path: Path) -> None:
        await module._action_write_paragraph({"path": str(doc_path), "text": "FontTest"})
        info = await module._action_list_paragraphs({"path": str(doc_path)})
        texts = [p["text"] for p in info["paragraphs"]]
        idx = texts.index("FontTest")

        result = await module._action_format_text(
            {
                "path": str(doc_path),
                "paragraph_index": idx,
                "run_index": 0,
                "font_name": "Courier New",
                "font_size": 12,
            }
        )
        assert result["formatted"] is True

    async def test_format_text_out_of_range_run_raises(
        self, module: WordModule, doc_path: Path
    ) -> None:
        # Paragraph 0 exists; request run 999 which does not
        with pytest.raises(IndexError):
            await module._action_format_text(
                {"path": str(doc_path), "paragraph_index": 0, "run_index": 999}
            )


@pytest.mark.integration
class TestSectionBreak:
    async def test_insert_section_break(self, module: WordModule, doc_path: Path) -> None:
        result = await module._action_insert_section_break(
            {"path": str(doc_path), "break_type": "nextPage"}
        )
        assert result["section_break_inserted"] is True
        assert result["break_type"] == "nextPage"

    async def test_insert_section_break_continuous(self, module: WordModule, doc_path: Path) -> None:
        result = await module._action_insert_section_break(
            {"path": str(doc_path), "break_type": "continuous"}
        )
        assert result["section_break_inserted"] is True


@pytest.mark.integration
class TestHyperlinkAndBookmark:
    async def test_insert_hyperlink(self, module: WordModule, doc_path: Path) -> None:
        result = await module._action_insert_hyperlink(
            {
                "path": str(doc_path),
                "paragraph_index": 0,
                "text": "Click here",
                "url": "https://example.com",
            }
        )
        assert result["inserted"] is True
        assert result["url"] == "https://example.com"
        assert result["text"] == "Click here"

    async def test_add_bookmark(self, module: WordModule, doc_path: Path) -> None:
        result = await module._action_add_bookmark(
            {"path": str(doc_path), "paragraph_index": 0, "name": "intro_section"}
        )
        assert result["added"] is True
        assert result["bookmark_name"] == "intro_section"

    async def test_add_comment(self, module: WordModule, doc_path: Path) -> None:
        result = await module._action_add_comment(
            {
                "path": str(doc_path),
                "paragraph_index": 0,
                "text": "This is a comment",
                "author": "TestAuthor",
            }
        )
        assert result["author"] == "TestAuthor"
        assert result["comment"] == "This is a comment"
        # Verify note was inserted as a paragraph
        content = await module._action_extract_text({"path": str(doc_path)})
        assert "TestAuthor" in content["text"]


@pytest.mark.integration
class TestTableOfContents:
    async def test_insert_toc(self, module: WordModule, doc_path: Path) -> None:
        # Add some heading paragraphs first
        await module._action_write_paragraph(
            {"path": str(doc_path), "text": "Chapter 1", "style": "Heading 1"}
        )
        result = await module._action_insert_toc(
            {"path": str(doc_path), "title": "Table of Contents", "max_depth": 2}
        )
        assert result["toc_inserted"] is True
        assert result["max_depth"] == 2

    async def test_insert_toc_no_title(self, module: WordModule, doc_path: Path) -> None:
        result = await module._action_insert_toc(
            {"path": str(doc_path), "title": "", "max_depth": 3}
        )
        assert result["toc_inserted"] is True


@pytest.mark.integration
class TestHeaderFooter:
    async def test_add_header(self, module: WordModule, doc_path: Path) -> None:
        result = await module._action_add_header_footer(
            {"path": str(doc_path), "header_text": "My Header", "alignment": "center"}
        )
        assert result["header_set"] is True
        assert result["footer_set"] is False

    async def test_add_footer(self, module: WordModule, doc_path: Path) -> None:
        result = await module._action_add_header_footer(
            {"path": str(doc_path), "footer_text": "My Footer"}
        )
        assert result["footer_set"] is True
        assert result["header_set"] is False

    async def test_add_footer_with_page_numbers(self, module: WordModule, doc_path: Path) -> None:
        result = await module._action_add_header_footer(
            {"path": str(doc_path), "page_numbers": True}
        )
        assert result["page_numbers"] is True
        assert result["footer_set"] is True

    async def test_header_footer_invalid_section_raises(
        self, module: WordModule, doc_path: Path
    ) -> None:
        with pytest.raises(IndexError):
            await module._action_add_header_footer(
                {"path": str(doc_path), "section": 999, "header_text": "X"}
            )


# ---------------------------------------------------------------------------
# Extended tests — uncovered branches
# ---------------------------------------------------------------------------


@pytest.mark.integration
class TestReadDocumentExtended:
    async def test_read_document_with_headers_footers(
        self, module: WordModule, doc_path: Path
    ) -> None:
        """read_document with include_headers_footers=True covers lines 377-390."""
        # Add a header first
        await module._action_add_header_footer(
            {"path": str(doc_path), "section": 0, "header_text": "My Header"}
        )
        result = await module._action_read_document(
            {"path": str(doc_path), "include_headers_footers": True}
        )
        assert "headers" in result
        assert "footers" in result
        assert isinstance(result["headers"], list)

    async def test_extract_text_with_tables(
        self, module: WordModule, doc_with_table: Path
    ) -> None:
        """extract_text with include_tables=True covers lines 449-455."""
        result = await module._action_extract_text(
            {"path": str(doc_with_table), "include_tables": True}
        )
        assert "text" in result
        assert "length" in result


@pytest.mark.integration
class TestWriteParagraphExtended:
    async def test_write_paragraph_insert_after(
        self, module: WordModule, doc_path: Path
    ) -> None:
        """write_paragraph with insert_after_paragraph covers lines 493-502."""
        await module._action_write_paragraph(
            {"path": str(doc_path), "text": "First paragraph"}
        )
        result = await module._action_write_paragraph(
            {
                "path": str(doc_path),
                "text": "Inserted after paragraph 0",
                "insert_after_paragraph": 0,
            }
        )
        assert "paragraph_index" in result

    async def test_write_paragraph_with_spacing(
        self, module: WordModule, doc_path: Path
    ) -> None:
        """write_paragraph with space_before, space_after, line_spacing covers lines 514-519."""
        result = await module._action_write_paragraph(
            {
                "path": str(doc_path),
                "text": "Spaced paragraph",
                "space_before": 6.0,
                "space_after": 6.0,
                "line_spacing": 1.5,
            }
        )
        assert "paragraph_index" in result


@pytest.mark.integration
class TestInsertListExtended:
    async def test_insert_list_with_indent_level(
        self, module: WordModule, doc_path: Path
    ) -> None:
        """insert_list with indent_level covers lines 662-669."""
        result = await module._action_insert_list(
            {
                "path": str(doc_path),
                "items": ["Sub item 1", "Sub item 2"],
                "list_style": "bullet",
                "indent_level": 1,
            }
        )
        assert result["items_inserted"] == 2

    async def test_insert_list_with_insert_after(
        self, module: WordModule, doc_path: Path
    ) -> None:
        """insert_list with insert_after_paragraph covers lines 676-681."""
        await module._action_write_paragraph(
            {"path": str(doc_path), "text": "Reference para"}
        )
        result = await module._action_insert_list(
            {
                "path": str(doc_path),
                "items": ["Item A", "Item B"],
                "list_style": "number",
                "insert_after_paragraph": 0,
            }
        )
        assert result["items_inserted"] == 2


@pytest.mark.integration
class TestInsertTableExtended:
    async def test_insert_table_with_header_row(
        self, module: WordModule, doc_path: Path
    ) -> None:
        """insert_table with has_header=True covers lines 711-719."""
        result = await module._action_insert_table(
            {
                "path": str(doc_path),
                "rows": 3,
                "cols": 3,
                "data": [["H1", "H2", "H3"], ["R1C1", "R1C2", "R1C3"]],
                "has_header": True,
            }
        )
        assert result["rows"] == 3

    async def test_insert_table_with_insert_after(
        self, module: WordModule, doc_path: Path
    ) -> None:
        """insert_table with insert_after_paragraph covers lines 721-723."""
        await module._action_write_paragraph(
            {"path": str(doc_path), "text": "Para before table"}
        )
        result = await module._action_insert_table(
            {
                "path": str(doc_path),
                "rows": 2,
                "cols": 2,
                "insert_after_paragraph": 0,
            }
        )
        assert result["rows"] == 2


@pytest.mark.integration
class TestModifyTableCellExtended:
    async def test_modify_cell_with_vertical_alignment_and_bg_color(
        self, module: WordModule, doc_with_table: Path
    ) -> None:
        """modify_table_cell with vertical_alignment and bg_color covers lines 763-778."""
        result = await module._action_modify_table_cell(
            {
                "path": str(doc_with_table),
                "table_index": 0,
                "row": 0,
                "col": 0,
                "text": "Styled",
                "alignment": "center",
                "vertical_alignment": "center",
                "bg_color": "FFFF00",
            }
        )
        assert "row" in result

    async def test_modify_cell_out_of_range_table_raises(
        self, module: WordModule, doc_with_table: Path
    ) -> None:
        """modify_table_cell table out-of-range covers line 746."""
        with pytest.raises(IndexError):
            await module._action_modify_table_cell(
                {
                    "path": str(doc_with_table),
                    "table_index": 99,
                    "row": 0,
                    "col": 0,
                }
            )


@pytest.mark.integration
class TestInsertImage:
    async def test_insert_image(self, module: WordModule, doc_path: Path, tmp_path: Path) -> None:
        """insert_image covers lines 819-849."""
        import struct
        import zlib

        # Build a minimal 1x1 white PNG
        def _make_tiny_png() -> bytes:
            sig = b"\x89PNG\r\n\x1a\n"
            ihdr_data = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
            ihdr_crc = struct.pack(">I", zlib.crc32(b"IHDR" + ihdr_data) & 0xFFFFFFFF)
            ihdr = struct.pack(">I", 13) + b"IHDR" + ihdr_data + ihdr_crc
            raw = b"\x00\xff\xff\xff"  # filter byte + RGB
            compressed = zlib.compress(raw)
            idat_crc = struct.pack(">I", zlib.crc32(b"IDAT" + compressed) & 0xFFFFFFFF)
            idat = struct.pack(">I", len(compressed)) + b"IDAT" + compressed + idat_crc
            iend_crc = struct.pack(">I", zlib.crc32(b"IEND") & 0xFFFFFFFF)
            iend = struct.pack(">I", 0) + b"IEND" + iend_crc
            return sig + ihdr + idat + iend

        img_path = tmp_path / "tiny.png"
        img_path.write_bytes(_make_tiny_png())

        result = await module._action_insert_image(
            {
                "path": str(doc_path),
                "image_path": str(img_path),
                "width_cm": 2.0,
                "alignment": "center",
            }
        )
        assert result["inserted"] is True
