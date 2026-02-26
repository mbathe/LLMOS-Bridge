"""Word module — python-docx-powered document automation.

All blocking I/O is delegated to ``asyncio.to_thread`` so the async event
loop is never blocked.  Heavy imports (docx) are deferred until first use.

Document instances are cached in ``self._doc_cache`` by their *source* path
so repeated operations on the same file avoid redundant disk reads.
"""

from __future__ import annotations

import asyncio
import subprocess
import threading
from pathlib import Path
from typing import Any

from llmos_bridge.modules.base import BaseModule, Platform
from llmos_bridge.modules.manifest import ActionSpec, ModuleManifest, ParamSpec
from llmos_bridge.protocol.params.word import (
    AddBookmarkParams,
    AddCommentParams,
    AddHeaderFooterParams,
    AddTableRowParams,
    ApplyStyleParams,
    CountWordsParams,
    CreateDocumentParams,
    DeleteParagraphParams,
    ExportToPdfParams,
    ExtractTextParams,
    FindReplaceParams,
    FormatTextParams,
    GetDocumentMetaParams,
    InsertHyperlinkParams,
    InsertImageParams,
    InsertListParams,
    InsertPageBreakParams,
    InsertSectionBreakParams,
    InsertTableOfContentsParams,
    InsertTableParams,
    ListParagraphsParams,
    ListTablesParams,
    ModifyTableCellParams,
    OpenDocumentParams,
    ReadDocumentParams,
    SaveDocumentParams,
    SetDefaultFontParams,
    SetDocumentPropertiesParams,
    SetMarginsParams,
    WriteParagraphParams,
)


class WordModule(BaseModule):
    """LLMOS Bridge module for creating and manipulating Word (.docx) documents."""

    MODULE_ID = "word"
    VERSION = "1.0.0"
    SUPPORTED_PLATFORMS = [Platform.ALL]

    def __init__(self) -> None:
        self._doc_cache: dict[str, Any] = {}
        self._path_locks: dict[str, threading.Lock] = {}
        self._meta_lock = threading.Lock()
        super().__init__()

    def _get_path_lock(self, path: str) -> threading.Lock:
        """Return (or create) a per-file threading.Lock for concurrent access control."""
        resolved = str(Path(path).resolve())
        with self._meta_lock:
            if resolved not in self._path_locks:
                self._path_locks[resolved] = threading.Lock()
            return self._path_locks[resolved]

    # ------------------------------------------------------------------
    # Dependency check
    # ------------------------------------------------------------------

    def _check_dependencies(self) -> None:
        try:
            import docx  # noqa: F401
        except ImportError as exc:
            from llmos_bridge.exceptions import ModuleLoadError

            raise ModuleLoadError(
                self.MODULE_ID,
                "python-docx is not installed. Run: pip install python-docx",
            ) from exc

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    def _get_doc(self, path: str) -> Any:
        """Return a cached Document or load from disk."""
        import docx

        if path not in self._doc_cache:
            self._doc_cache[path] = docx.Document(path)
        return self._doc_cache[path]

    def _save_doc(self, path: str, doc: Any, output_path: str | None = None) -> str:
        """Save *doc* and update the cache.  Returns the actual save path."""
        save_path = output_path or path
        doc.save(save_path)
        # Keep cache in sync — if saved to a new path, cache under that path too.
        self._doc_cache[save_path] = doc
        if output_path and output_path != path:
            # Original path still holds the old version; evict stale entry.
            self._doc_cache.pop(path, None)
        return save_path

    def _get_para(self, doc: Any, index: int) -> Any:
        """Return ``doc.paragraphs[index]``, raising IndexError on bounds failure."""
        paragraphs = doc.paragraphs
        if index < 0 or index >= len(paragraphs):
            raise IndexError(
                f"Paragraph index {index} is out of range "
                f"(document has {len(paragraphs)} paragraphs)."
            )
        return paragraphs[index]

    def _apply_run_format(self, run: Any, p: Any) -> None:
        """Apply character-level formatting fields from *p* onto *run*."""
        from docx.shared import Pt, RGBColor

        if getattr(p, "bold", None) is not None:
            run.bold = p.bold
        if getattr(p, "italic", None) is not None:
            run.italic = p.italic
        if getattr(p, "underline", None) is not None:
            run.underline = p.underline
        if getattr(p, "strikethrough", None) is not None:
            run.font.strike = p.strikethrough
        if getattr(p, "all_caps", None) is not None:
            run.font.all_caps = p.all_caps
        if getattr(p, "font_name", None):
            run.font.name = p.font_name
        if getattr(p, "font_size", None):
            run.font.size = Pt(p.font_size)
        if getattr(p, "font_color", None):
            hex_color = p.font_color.lstrip("#")
            run.font.color.rgb = RGBColor(
                int(hex_color[0:2], 16),
                int(hex_color[2:4], 16),
                int(hex_color[4:6], 16),
            )
        if getattr(p, "highlight_color", None) and p.highlight_color != "none":
            from docx.enum.text import WD_COLOR_INDEX

            color_map = {
                "yellow": WD_COLOR_INDEX.YELLOW,
                "green": WD_COLOR_INDEX.GREEN,
                "cyan": WD_COLOR_INDEX.CYAN,
                "magenta": WD_COLOR_INDEX.MAGENTA,
                "blue": WD_COLOR_INDEX.BLUE,
                "red": WD_COLOR_INDEX.RED,
                "darkBlue": WD_COLOR_INDEX.DARK_BLUE,
                "darkCyan": WD_COLOR_INDEX.DARK_CYAN,
                "darkGreen": WD_COLOR_INDEX.DARK_GREEN,
                "darkMagenta": WD_COLOR_INDEX.DARK_MAGENTA,
                "darkRed": WD_COLOR_INDEX.DARK_RED,
                "darkYellow": WD_COLOR_INDEX.DARK_YELLOW,
                "darkGray": WD_COLOR_INDEX.DARK_GRAY,
                "lightGray": WD_COLOR_INDEX.LIGHT_GRAY,
                "white": WD_COLOR_INDEX.WHITE,
                "black": WD_COLOR_INDEX.BLACK,
            }
            hi = color_map.get(p.highlight_color)
            if hi is not None:
                run.font.highlight_color = hi

    def _set_paragraph_alignment(self, paragraph: Any, alignment: str) -> None:
        """Set paragraph alignment from a string literal."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        mapping = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        if alignment in mapping:
            paragraph.alignment = mapping[alignment]

    def _iter_paragraphs_and_cells(self, doc: Any):
        """Yield every paragraph in the document, including those in table cells."""
        yield from doc.paragraphs
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs

    # ------------------------------------------------------------------
    # Lifecycle actions
    # ------------------------------------------------------------------

    async def _action_open_document(self, params: dict[str, Any]) -> dict[str, Any]:
        p = OpenDocumentParams.model_validate(params)

        def _open() -> dict[str, Any]:
            with self._get_path_lock(p.path):
                doc = self._get_doc(p.path)
                return {
                    "path": p.path,
                    "paragraph_count": len(doc.paragraphs),
                    "table_count": len(doc.tables),
                    "section_count": len(doc.sections),
                }

        return await asyncio.to_thread(_open)

    async def _action_create_document(self, params: dict[str, Any]) -> dict[str, Any]:
        p = CreateDocumentParams.model_validate(params)

        def _create() -> dict[str, Any]:
            with self._get_path_lock(p.output_path):
                import docx
                from docx.shared import Pt

                doc = docx.Document()
                props = doc.core_properties
                if p.title is not None:
                    props.title = p.title
                if p.author is not None:
                    props.author = p.author
                if p.subject is not None:
                    props.subject = p.subject
                if p.description is not None:
                    props.comments = p.description

                if p.default_font or p.default_font_size:
                    normal = doc.styles["Normal"]
                    if p.default_font:
                        normal.font.name = p.default_font
                    if p.default_font_size:
                        normal.font.size = Pt(p.default_font_size)

                out = Path(p.output_path)
                out.parent.mkdir(parents=True, exist_ok=True)
                doc.save(str(out))
                self._doc_cache[str(out)] = doc
                return {
                    "path": str(out),
                    "created": True,
                }

        return await asyncio.to_thread(_create)

    async def _action_save_document(self, params: dict[str, Any]) -> dict[str, Any]:
        p = SaveDocumentParams.model_validate(params)

        def _save() -> dict[str, Any]:
            with self._get_path_lock(p.path):
                doc = self._get_doc(p.path)
                saved_path = self._save_doc(p.path, doc, p.output_path)
                return {"path": saved_path, "saved": True}

        return await asyncio.to_thread(_save)

    async def _action_set_document_properties(self, params: dict[str, Any]) -> dict[str, Any]:
        p = SetDocumentPropertiesParams.model_validate(params)

        def _set_props() -> dict[str, Any]:
            with self._get_path_lock(p.path):
                doc = self._get_doc(p.path)
                props = doc.core_properties
                if p.title is not None:
                    props.title = p.title
                if p.author is not None:
                    props.author = p.author
                if p.subject is not None:
                    props.subject = p.subject
                if p.description is not None:
                    props.comments = p.description
                if p.keywords is not None:
                    props.keywords = p.keywords
                if p.category is not None:
                    props.category = p.category
                if p.company is not None:
                    props.company = p.company
                if p.language is not None:
                    props.language = p.language
                self._save_doc(p.path, doc)
                return {"path": p.path, "updated": True}

        return await asyncio.to_thread(_set_props)

    async def _action_get_document_meta(self, params: dict[str, Any]) -> dict[str, Any]:
        p = GetDocumentMetaParams.model_validate(params)

        def _get_meta() -> dict[str, Any]:
            with self._get_path_lock(p.path):
                doc = self._get_doc(p.path)
                props = doc.core_properties

                def _dt(val: Any) -> str | None:
                    if val is None:
                        return None
                    try:
                        return val.isoformat()
                    except AttributeError:
                        return str(val)

                return {
                    "path": p.path,
                    "title": props.title,
                    "author": props.author,
                    "subject": props.subject,
                    "description": props.comments,
                    "keywords": props.keywords,
                    "category": props.category,
                    "language": props.language,
                    "created": _dt(props.created),
                    "modified": _dt(props.modified),
                    "last_modified_by": props.last_modified_by,
                    "revision": props.revision,
                    "version": props.version,
                }

        return await asyncio.to_thread(_get_meta)

    async def _action_set_margins(self, params: dict[str, Any]) -> dict[str, Any]:
        p = SetMarginsParams.model_validate(params)

        def _set_margins() -> dict[str, Any]:
            with self._get_path_lock(p.path):
                from docx.shared import Cm

                doc = self._get_doc(p.path)
                sections = doc.sections
                if p.section >= len(sections):
                    raise IndexError(
                        f"Section index {p.section} is out of range "
                        f"(document has {len(sections)} section(s))."
                    )
                sec = sections[p.section]
                sec.top_margin = Cm(p.top)
                sec.bottom_margin = Cm(p.bottom)
                sec.left_margin = Cm(p.left)
                sec.right_margin = Cm(p.right)
                self._save_doc(p.path, doc)
                return {
                    "path": p.path,
                    "section": p.section,
                    "top_cm": p.top,
                    "bottom_cm": p.bottom,
                    "left_cm": p.left,
                    "right_cm": p.right,
                }

        return await asyncio.to_thread(_set_margins)

    async def _action_set_default_font(self, params: dict[str, Any]) -> dict[str, Any]:
        p = SetDefaultFontParams.model_validate(params)

        def _set_font() -> dict[str, Any]:
            with self._get_path_lock(p.path):
                from docx.shared import Pt

                doc = self._get_doc(p.path)
                normal = doc.styles["Normal"]
                normal.font.name = p.font_name
                if p.font_size:
                    normal.font.size = Pt(p.font_size)
                self._save_doc(p.path, doc)
                return {"path": p.path, "font_name": p.font_name, "font_size": p.font_size}

        return await asyncio.to_thread(_set_font)

    # ------------------------------------------------------------------
    # Read operations
    # ------------------------------------------------------------------

    async def _action_read_document(self, params: dict[str, Any]) -> dict[str, Any]:
        p = ReadDocumentParams.model_validate(params)

        def _read() -> dict[str, Any]:
            with self._get_path_lock(p.path):
                doc = self._get_doc(p.path)
                paragraphs = [
                    {"text": para.text, "style": para.style.name}
                    for para in doc.paragraphs
                ]
                result: dict[str, Any] = {"path": p.path, "paragraphs": paragraphs}

                if p.include_tables:
                    tables = []
                    for t_idx, table in enumerate(doc.tables):
                        rows_data = []
                        for row in table.rows:
                            rows_data.append([cell.text for cell in row.cells])
                        tables.append({"index": t_idx, "data": rows_data})
                    result["tables"] = tables

                if p.include_headers_footers:
                    headers = []
                    footers = []
                    for s_idx, section in enumerate(doc.sections):
                        hdr_text = "\n".join(
                            p.text for p in section.header.paragraphs
                        )
                        ftr_text = "\n".join(
                            p.text for p in section.footer.paragraphs
                        )
                        headers.append({"section": s_idx, "text": hdr_text})
                        footers.append({"section": s_idx, "text": ftr_text})
                    result["headers"] = headers
                    result["footers"] = footers

                return result

        return await asyncio.to_thread(_read)

    async def _action_list_paragraphs(self, params: dict[str, Any]) -> dict[str, Any]:
        p = ListParagraphsParams.model_validate(params)

        def _list() -> dict[str, Any]:
            with self._get_path_lock(p.path):
                doc = self._get_doc(p.path)
                result = []
                for idx, para in enumerate(doc.paragraphs):
                    is_empty = not para.text.strip()
                    if not p.include_empty and is_empty:
                        continue
                    if p.style_filter and para.style.name != p.style_filter:
                        continue
                    result.append(
                        {
                            "index": idx,
                            "text": para.text,
                            "style": para.style.name,
                            "is_empty": is_empty,
                        }
                    )
                return {"path": p.path, "paragraphs": result, "count": len(result)}

        return await asyncio.to_thread(_list)

    async def _action_list_tables(self, params: dict[str, Any]) -> dict[str, Any]:
        p = ListTablesParams.model_validate(params)

        def _list() -> dict[str, Any]:
            with self._get_path_lock(p.path):
                doc = self._get_doc(p.path)
                tables = []
                for idx, table in enumerate(doc.tables):
                    rows = len(table.rows)
                    cols = len(table.columns) if table.rows else 0
                    first_cell = table.cell(0, 0).text if rows > 0 and cols > 0 else ""
                    tables.append(
                        {
                            "index": idx,
                            "rows": rows,
                            "cols": cols,
                            "first_cell_text": first_cell,
                        }
                    )
                return {"path": p.path, "tables": tables, "count": len(tables)}

        return await asyncio.to_thread(_list)

    async def _action_extract_text(self, params: dict[str, Any]) -> dict[str, Any]:
        p = ExtractTextParams.model_validate(params)

        def _extract() -> dict[str, Any]:
            with self._get_path_lock(p.path):
                doc = self._get_doc(p.path)
                texts = [para.text for para in doc.paragraphs]

                if p.include_tables:
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for cell_para in cell.paragraphs:
                                    if cell_para.text:
                                        texts.append(cell_para.text)

                full_text = p.separator.join(texts)
                return {
                    "path": p.path,
                    "text": full_text,
                    "length": len(full_text),
                }

        return await asyncio.to_thread(_extract)

    async def _action_count_words(self, params: dict[str, Any]) -> dict[str, Any]:
        p = CountWordsParams.model_validate(params)

        def _count() -> dict[str, Any]:
            with self._get_path_lock(p.path):
                doc = self._get_doc(p.path)
                word_count = 0
                for para in self._iter_paragraphs_and_cells(doc):
                    word_count += len(para.text.split())
                return {"path": p.path, "word_count": word_count}

        return await asyncio.to_thread(_count)

    # ------------------------------------------------------------------
    # Paragraph operations
    # ------------------------------------------------------------------

    async def _action_write_paragraph(self, params: dict[str, Any]) -> dict[str, Any]:
        p = WriteParagraphParams.model_validate(params)

        def _write() -> dict[str, Any]:
            with self._get_path_lock(p.path):
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.oxml.ns import qn
                from docx.shared import Pt
                from lxml import etree

                doc = self._get_doc(p.path)

                if p.insert_after_paragraph is not None:
                    # Insert a new paragraph after the specified index.
                    ref_para = self._get_para(doc, p.insert_after_paragraph)
                    # Build a new paragraph element by adding it to the document body
                    # then moving it after the reference paragraph.
                    new_para = doc.add_paragraph(p.text, style=p.style)
                    # Move the newly appended element to after the reference paragraph.
                    ref_para._element.addnext(new_para._element)
                    # The paragraph was appended last but is now at the correct position.
                    paragraph = new_para
                else:
                    paragraph = doc.add_paragraph(p.text, style=p.style)

                # Apply run-level formatting to the run created by add_paragraph.
                if paragraph.runs:
                    run = paragraph.runs[0]
                    self._apply_run_format(run, p)

                # Paragraph-level formatting.
                self._set_paragraph_alignment(paragraph, p.alignment)

                if p.space_before is not None:
                    paragraph.paragraph_format.space_before = Pt(p.space_before)
                if p.space_after is not None:
                    paragraph.paragraph_format.space_after = Pt(p.space_after)
                if p.line_spacing is not None:
                    paragraph.paragraph_format.line_spacing = p.line_spacing

                self._save_doc(p.path, doc)
                para_elements = [para._p for para in doc.paragraphs]
                try:
                    para_idx = para_elements.index(paragraph._p)
                except ValueError:
                    para_idx = len(doc.paragraphs) - 1
                return {
                    "path": p.path,
                    "paragraph_index": para_idx,
                    "text": paragraph.text,
                    "style": paragraph.style.name,
                }

        return await asyncio.to_thread(_write)

    async def _action_format_text(self, params: dict[str, Any]) -> dict[str, Any]:
        p = FormatTextParams.model_validate(params)

        def _format() -> dict[str, Any]:
            with self._get_path_lock(p.path):
                doc = self._get_doc(p.path)
                paragraph = self._get_para(doc, p.paragraph_index)
                if p.run_index >= len(paragraph.runs):
                    raise IndexError(
                        f"Run index {p.run_index} is out of range "
                        f"(paragraph has {len(paragraph.runs)} run(s))."
                    )
                run = paragraph.runs[p.run_index]
                self._apply_run_format(run, p)
                self._save_doc(p.path, doc)
                return {
                    "path": p.path,
                    "paragraph_index": p.paragraph_index,
                    "run_index": p.run_index,
                    "formatted": True,
                }

        return await asyncio.to_thread(_format)

    async def _action_apply_style(self, params: dict[str, Any]) -> dict[str, Any]:
        p = ApplyStyleParams.model_validate(params)

        def _apply() -> dict[str, Any]:
            with self._get_path_lock(p.path):
                doc = self._get_doc(p.path)
                paragraph = self._get_para(doc, p.paragraph_index)
                paragraph.style = p.style
                self._save_doc(p.path, doc)
                return {
                    "path": p.path,
                    "paragraph_index": p.paragraph_index,
                    "style": p.style,
                }

        return await asyncio.to_thread(_apply)

    async def _action_delete_paragraph(self, params: dict[str, Any]) -> dict[str, Any]:
        p = DeleteParagraphParams.model_validate(params)

        def _delete() -> dict[str, Any]:
            with self._get_path_lock(p.path):
                doc = self._get_doc(p.path)
                paragraph = self._get_para(doc, p.paragraph_index)
                parent = paragraph._p.getparent()
                parent.remove(paragraph._p)
                self._save_doc(p.path, doc)
                return {
                    "path": p.path,
                    "deleted_index": p.paragraph_index,
                    "remaining_paragraphs": len(doc.paragraphs),
                }

        return await asyncio.to_thread(_delete)

    async def _action_insert_page_break(self, params: dict[str, Any]) -> dict[str, Any]:
        p = InsertPageBreakParams.model_validate(params)

        def _insert() -> dict[str, Any]:
            with self._get_path_lock(p.path):
                from docx.enum.text import WD_BREAK

                doc = self._get_doc(p.path)
                new_para = doc.add_paragraph()
                run = new_para.add_run()
                run.add_break(WD_BREAK.PAGE)

                if p.after_paragraph is not None:
                    ref_para = self._get_para(doc, p.after_paragraph)
                    ref_para._element.addnext(new_para._element)

                self._save_doc(p.path, doc)
                return {"path": p.path, "page_break_inserted": True}

        return await asyncio.to_thread(_insert)

    async def _action_insert_section_break(self, params: dict[str, Any]) -> dict[str, Any]:
        p = InsertSectionBreakParams.model_validate(params)

        def _insert() -> dict[str, Any]:
            with self._get_path_lock(p.path):
                from docx.enum.section import WD_SECTION
                from docx.oxml.ns import qn
                from lxml import etree

                doc = self._get_doc(p.path)

                break_type_map = {
                    "nextPage": WD_SECTION.NEW_PAGE,
                    "continuous": WD_SECTION.CONTINUOUS,
                    "evenPage": WD_SECTION.EVEN_PAGE,
                    "oddPage": WD_SECTION.ODD_PAGE,
                }

                # Add a new section with the given break type.
                new_section = doc.add_section(break_type_map.get(p.break_type, WD_SECTION.NEW_PAGE))

                if p.after_paragraph is not None:
                    # python-docx adds sections at the end; we need to move the
                    # sectPr element to after the target paragraph's pPr.
                    ref_para = self._get_para(doc, p.after_paragraph)
                    sect_pr = new_section._sectPr
                    ref_para._p.append(sect_pr)

                self._save_doc(p.path, doc)
                return {
                    "path": p.path,
                    "break_type": p.break_type,
                    "section_break_inserted": True,
                }

        return await asyncio.to_thread(_insert)

    async def _action_insert_list(self, params: dict[str, Any]) -> dict[str, Any]:
        p = InsertListParams.model_validate(params)

        def _insert() -> dict[str, Any]:
            with self._get_path_lock(p.path):
                doc = self._get_doc(p.path)

                style_map = {
                    "bullet": "List Bullet",
                    "number": "List Number",
                    "alpha": "List Number",   # python-docx does not have separate alpha/roman
                    "roman": "List Number",
                }
                style_name = style_map.get(p.list_style, "List Bullet")
                # Adjust indent level by appending the level number to the style name.
                if p.indent_level > 0:
                    candidate = f"{style_name} {p.indent_level + 1}"
                    # Only use the levelled style if it exists in the document.
                    try:
                        doc.styles[candidate]
                        style_name = candidate
                    except KeyError:
                        pass

                inserted: list[dict[str, Any]] = []
                prev_element = None

                for idx, item_text in enumerate(p.items):
                    new_para = doc.add_paragraph(item_text, style=style_name)
                    if p.insert_after_paragraph is not None:
                        if prev_element is None:
                            ref_para = self._get_para(doc, p.insert_after_paragraph)
                            ref_para._element.addnext(new_para._element)
                        else:
                            prev_element.addnext(new_para._element)
                    prev_element = new_para._element
                    inserted.append({"index": idx, "text": item_text})

                self._save_doc(p.path, doc)
                return {
                    "path": p.path,
                    "items_inserted": len(inserted),
                    "style": style_name,
                }

        return await asyncio.to_thread(_insert)

    # ------------------------------------------------------------------
    # Table operations
    # ------------------------------------------------------------------

    async def _action_insert_table(self, params: dict[str, Any]) -> dict[str, Any]:
        p = InsertTableParams.model_validate(params)

        def _insert() -> dict[str, Any]:
            with self._get_path_lock(p.path):
                doc = self._get_doc(p.path)
                table = doc.add_table(rows=p.rows, cols=p.cols)
                table.style = p.style

                if p.data:
                    for r_idx, row_data in enumerate(p.data[: p.rows]):
                        for c_idx, cell_text in enumerate(row_data[: p.cols]):
                            table.cell(r_idx, c_idx).text = str(cell_text)

                if p.has_header and p.rows > 0:
                    from docx.oxml.ns import qn
                    from docx.oxml import OxmlElement

                    header_row = table.rows[0]
                    tr = header_row._tr
                    trPr = tr.get_or_add_trPr()
                    tblHeader = OxmlElement("w:tblHeader")
                    trPr.append(tblHeader)

                if p.insert_after_paragraph is not None:
                    ref_para = self._get_para(doc, p.insert_after_paragraph)
                    ref_para._element.addnext(table._tbl)

                self._save_doc(p.path, doc)
                return {
                    "path": p.path,
                    "table_index": len(doc.tables) - 1,
                    "rows": p.rows,
                    "cols": p.cols,
                }

        return await asyncio.to_thread(_insert)

    async def _action_modify_table_cell(self, params: dict[str, Any]) -> dict[str, Any]:
        p = ModifyTableCellParams.model_validate(params)

        def _modify() -> dict[str, Any]:
            with self._get_path_lock(p.path):
                from docx.enum.table import WD_ALIGN_VERTICAL
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
                from docx.shared import RGBColor

                doc = self._get_doc(p.path)
                if p.table_index >= len(doc.tables):
                    raise IndexError(
                        f"Table index {p.table_index} is out of range "
                        f"(document has {len(doc.tables)} table(s))."
                    )
                table = doc.tables[p.table_index]
                cell = table.cell(p.row, p.col)

                if p.text is not None:
                    cell.text = p.text

                # Apply run-level formatting to all runs in the cell paragraphs.
                for para in cell.paragraphs:
                    for run in para.runs:
                        self._apply_run_format(run, p)
                    if p.alignment:
                        self._set_paragraph_alignment(para, p.alignment)

                if p.vertical_alignment:
                    va_map = {
                        "top": WD_ALIGN_VERTICAL.TOP,
                        "center": WD_ALIGN_VERTICAL.CENTER,
                        "bottom": WD_ALIGN_VERTICAL.BOTTOM,
                    }
                    cell.vertical_alignment = va_map.get(p.vertical_alignment)

                if p.bg_color:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"), p.bg_color.lstrip("#").upper())
                    tcPr.append(shd)

                self._save_doc(p.path, doc)
                return {
                    "path": p.path,
                    "table_index": p.table_index,
                    "row": p.row,
                    "col": p.col,
                    "updated": True,
                }

        return await asyncio.to_thread(_modify)

    async def _action_add_table_row(self, params: dict[str, Any]) -> dict[str, Any]:
        p = AddTableRowParams.model_validate(params)

        def _add_row() -> dict[str, Any]:
            with self._get_path_lock(p.path):
                doc = self._get_doc(p.path)
                if p.table_index >= len(doc.tables):
                    raise IndexError(
                        f"Table index {p.table_index} is out of range "
                        f"(document has {len(doc.tables)} table(s))."
                    )
                table = doc.tables[p.table_index]
                new_row = table.add_row()
                for c_idx, cell_text in enumerate(p.data[: len(new_row.cells)]):
                    new_row.cells[c_idx].text = str(cell_text)

                self._save_doc(p.path, doc)
                return {
                    "path": p.path,
                    "table_index": p.table_index,
                    "new_row_index": len(table.rows) - 1,
                }

        return await asyncio.to_thread(_add_row)

    # ------------------------------------------------------------------
    # Rich content
    # ------------------------------------------------------------------

    async def _action_insert_image(self, params: dict[str, Any]) -> dict[str, Any]:
        p = InsertImageParams.model_validate(params)

        def _insert() -> dict[str, Any]:
            with self._get_path_lock(p.path):
                from docx.shared import Cm

                doc = self._get_doc(p.path)
                img_para = doc.add_paragraph()
                self._set_paragraph_alignment(img_para, p.alignment)

                run = img_para.add_run()
                kwargs: dict[str, Any] = {"image_path_or_stream": p.image_path}
                if p.width_cm is not None:
                    kwargs["width"] = Cm(p.width_cm)
                if p.height_cm is not None:
                    kwargs["height"] = Cm(p.height_cm)
                run.add_picture(**kwargs)

                if p.insert_after_paragraph is not None:
                    ref_para = self._get_para(doc, p.insert_after_paragraph)
                    ref_para._element.addnext(img_para._element)

                if p.caption:
                    caption_para = doc.add_paragraph(p.caption, style="Caption")
                    self._set_paragraph_alignment(caption_para, p.alignment)
                    img_para._element.addnext(caption_para._element)

                self._save_doc(p.path, doc)
                return {"path": p.path, "image_path": p.image_path, "inserted": True}

        return await asyncio.to_thread(_insert)

    async def _action_insert_hyperlink(self, params: dict[str, Any]) -> dict[str, Any]:
        p = InsertHyperlinkParams.model_validate(params)

        def _insert() -> dict[str, Any]:
            with self._get_path_lock(p.path):
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
                from docx.shared import Pt

                doc = self._get_doc(p.path)
                paragraph = self._get_para(doc, p.paragraph_index)

                # Add the relationship and get its rId.
                part = doc.part
                r_id = part.relate_to(
                    p.url,
                    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                    is_external=True,
                )

                # Build <w:hyperlink r:id="rId..." w:history="1">
                hyperlink = OxmlElement("w:hyperlink")
                hyperlink.set(qn("r:id"), r_id)
                hyperlink.set(qn("w:history"), "1")

                # Build <w:r> (run) inside the hyperlink.
                new_run = OxmlElement("w:r")
                rPr = OxmlElement("w:rPr")

                # Apply Hyperlink character style via <w:rStyle>.
                rStyle = OxmlElement("w:rStyle")
                rStyle.set(qn("w:val"), "Hyperlink")
                rPr.append(rStyle)

                if p.font_name:
                    rFonts = OxmlElement("w:rFonts")
                    rFonts.set(qn("w:ascii"), p.font_name)
                    rFonts.set(qn("w:hAnsi"), p.font_name)
                    rPr.append(rFonts)

                if p.font_size:
                    sz = OxmlElement("w:sz")
                    sz.set(qn("w:val"), str(p.font_size * 2))
                    rPr.append(sz)

                new_run.append(rPr)

                # <w:t> with the link text.
                t = OxmlElement("w:t")
                t.text = p.text
                new_run.append(t)
                hyperlink.append(new_run)

                paragraph._p.append(hyperlink)

                self._save_doc(p.path, doc)
                return {
                    "path": p.path,
                    "paragraph_index": p.paragraph_index,
                    "url": p.url,
                    "text": p.text,
                    "inserted": True,
                }

        return await asyncio.to_thread(_insert)

    async def _action_add_bookmark(self, params: dict[str, Any]) -> dict[str, Any]:
        p = AddBookmarkParams.model_validate(params)

        def _add() -> dict[str, Any]:
            with self._get_path_lock(p.path):
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn

                doc = self._get_doc(p.path)
                paragraph = self._get_para(doc, p.paragraph_index)

                # Use a simple incrementing id for bookmark uniqueness.
                bookmark_id = str(abs(hash(p.name)) % 100000)

                bookmark_start = OxmlElement("w:bookmarkStart")
                bookmark_start.set(qn("w:id"), bookmark_id)
                bookmark_start.set(qn("w:name"), p.name)

                bookmark_end = OxmlElement("w:bookmarkEnd")
                bookmark_end.set(qn("w:id"), bookmark_id)

                # Wrap existing content between bookmarkStart/End.
                p_elem = paragraph._p
                # Insert bookmarkStart before the first run (or at the beginning).
                first_run = p_elem.find(qn("w:r"))
                if first_run is not None:
                    first_run.addprevious(bookmark_start)
                else:
                    p_elem.append(bookmark_start)
                p_elem.append(bookmark_end)

                self._save_doc(p.path, doc)
                return {
                    "path": p.path,
                    "paragraph_index": p.paragraph_index,
                    "bookmark_name": p.name,
                    "added": True,
                }

        return await asyncio.to_thread(_add)

    async def _action_add_comment(self, params: dict[str, Any]) -> dict[str, Any]:
        p = AddCommentParams.model_validate(params)

        def _add() -> dict[str, Any]:
            with self._get_path_lock(p.path):
                doc = self._get_doc(p.path)
                paragraph = self._get_para(doc, p.paragraph_index)

                # python-docx does not support native comment XML out of the box.
                # We append a visible note paragraph immediately after the target
                # paragraph as a reliable fallback.
                note_text = f"[COMMENT by {p.author}: {p.text}]"
                note_para = doc.add_paragraph(note_text, style="Normal")
                paragraph._element.addnext(note_para._element)

                self._save_doc(p.path, doc)
                return {
                    "path": p.path,
                    "paragraph_index": p.paragraph_index,
                    "author": p.author,
                    "comment": p.text,
                    "note": "Comment added as a visible note paragraph (python-docx limitation).",
                }

        return await asyncio.to_thread(_add)

    async def _action_insert_toc(self, params: dict[str, Any]) -> dict[str, Any]:
        p = InsertTableOfContentsParams.model_validate(params)

        def _insert() -> dict[str, Any]:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            doc = self._get_doc(p.path)

            # Heading paragraph for the TOC.
            if p.title:
                title_para = doc.add_paragraph(p.title, style="TOC Heading")
            else:
                title_para = None

            # TOC paragraph with field code.
            toc_para = doc.add_paragraph()
            run = toc_para.add_run()

            fld_char_begin = OxmlElement("w:fldChar")
            fld_char_begin.set(qn("w:fldCharType"), "begin")

            instr_text = OxmlElement("w:instrText")
            instr_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            instr_text.text = f" TOC \\o '1-{p.max_depth}' \\h \\z \\u "

            fld_char_separate = OxmlElement("w:fldChar")
            fld_char_separate.set(qn("w:fldCharType"), "separate")

            fld_char_end = OxmlElement("w:fldChar")
            fld_char_end.set(qn("w:fldCharType"), "end")

            run._r.append(fld_char_begin)
            run._r.append(instr_text)
            run._r.append(fld_char_separate)
            run._r.append(fld_char_end)

            if p.insert_after_paragraph is not None:
                ref_para = self._get_para(doc, p.insert_after_paragraph)
                if title_para is not None:
                    ref_para._element.addnext(title_para._element)
                    title_para._element.addnext(toc_para._element)
                else:
                    ref_para._element.addnext(toc_para._element)

            self._save_doc(p.path, doc)
            return {
                "path": p.path,
                "toc_inserted": True,
                "max_depth": p.max_depth,
                "note": "Open the document in Word and press Ctrl+A then F9 to update the TOC.",
            }

        return await asyncio.to_thread(_insert)

    # ------------------------------------------------------------------
    # Header / footer
    # ------------------------------------------------------------------

    async def _action_add_header_footer(self, params: dict[str, Any]) -> dict[str, Any]:
        p = AddHeaderFooterParams.model_validate(params)

        def _add() -> dict[str, Any]:
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            doc = self._get_doc(p.path)
            if p.section >= len(doc.sections):
                raise IndexError(
                    f"Section index {p.section} is out of range "
                    f"(document has {len(doc.sections)} section(s))."
                )
            section = doc.sections[p.section]

            if p.header_text is not None:
                section.header.is_linked_to_previous = False
                hdr_para = section.header.paragraphs[0]
                hdr_para.clear()
                run = hdr_para.add_run(p.header_text)
                self._set_paragraph_alignment(hdr_para, p.alignment)

            if p.footer_text is not None or p.page_numbers:
                section.footer.is_linked_to_previous = False
                ftr_para = section.footer.paragraphs[0]
                ftr_para.clear()
                if p.footer_text:
                    ftr_para.add_run(p.footer_text)
                if p.page_numbers:
                    # Insert a PAGE field for automatic page numbering.
                    ftr_para.add_run("  ")
                    fld_run = ftr_para.add_run()
                    fld_char_begin = OxmlElement("w:fldChar")
                    fld_char_begin.set(qn("w:fldCharType"), "begin")
                    instr = OxmlElement("w:instrText")
                    instr.text = " PAGE "
                    fld_char_end = OxmlElement("w:fldChar")
                    fld_char_end.set(qn("w:fldCharType"), "end")
                    fld_run._r.append(fld_char_begin)
                    fld_run._r.append(instr)
                    fld_run._r.append(fld_char_end)
                self._set_paragraph_alignment(ftr_para, p.alignment)

            self._save_doc(p.path, doc)
            return {
                "path": p.path,
                "section": p.section,
                "header_set": p.header_text is not None,
                "footer_set": p.footer_text is not None or p.page_numbers,
                "page_numbers": p.page_numbers,
            }

        return await asyncio.to_thread(_add)

    # ------------------------------------------------------------------
    # Search & replace
    # ------------------------------------------------------------------

    async def _action_find_replace(self, params: dict[str, Any]) -> dict[str, Any]:
        p = FindReplaceParams.model_validate(params)

        def _find_replace() -> dict[str, Any]:
            import re

            doc = self._get_doc(p.path)

            # Build the pattern.
            flags = 0 if p.case_sensitive else re.IGNORECASE
            find_text = re.escape(p.find) if not p.whole_word else rf"\b{re.escape(p.find)}\b"
            pattern = re.compile(find_text, flags)

            replacements = 0

            for para in self._iter_paragraphs_and_cells(doc):
                # Replace within individual runs where possible.
                # For cross-run matches, reassemble the full paragraph text,
                # replace, then rewrite as a single run preserving the first
                # run's formatting.
                full_text = para.text
                if not pattern.search(full_text):
                    continue

                new_text, n = pattern.subn(p.replace, full_text)
                replacements += n

                if n > 0:
                    # Preserve formatting of the first run if it exists.
                    first_run_fmt: dict[str, Any] = {}
                    if para.runs:
                        fr = para.runs[0]
                        first_run_fmt = {
                            "bold": fr.bold,
                            "italic": fr.italic,
                            "underline": fr.underline,
                            "font_name": fr.font.name,
                            "font_size": fr.font.size,
                        }

                    # Clear all runs.
                    for run in para.runs:
                        run.text = ""

                    # Set text on the first run (or add one).
                    if para.runs:
                        para.runs[0].text = new_text
                        run = para.runs[0]
                        if first_run_fmt.get("bold") is not None:
                            run.bold = first_run_fmt["bold"]
                        if first_run_fmt.get("italic") is not None:
                            run.italic = first_run_fmt["italic"]
                        if first_run_fmt.get("underline") is not None:
                            run.underline = first_run_fmt["underline"]
                    else:
                        para.add_run(new_text)

            self._save_doc(p.path, doc)
            return {
                "path": p.path,
                "find": p.find,
                "replace": p.replace,
                "replacements_made": replacements,
            }

        return await asyncio.to_thread(_find_replace)

    # ------------------------------------------------------------------
    # Export
    # ------------------------------------------------------------------

    async def _action_export_to_pdf(self, params: dict[str, Any]) -> dict[str, Any]:
        p = ExportToPdfParams.model_validate(params)

        def _export() -> dict[str, Any]:
            docx_path = Path(p.path)
            pdf_path = Path(p.output_path)
            pdf_path.parent.mkdir(parents=True, exist_ok=True)

            if p.use_libreoffice:
                result = subprocess.run(
                    [
                        "libreoffice",
                        "--headless",
                        "--convert-to",
                        "pdf",
                        "--outdir",
                        str(pdf_path.parent),
                        str(docx_path),
                    ],
                    capture_output=True,
                    text=True,
                    timeout=120,
                )
                if result.returncode != 0:
                    raise RuntimeError(
                        f"LibreOffice conversion failed (exit {result.returncode}): "
                        f"{result.stderr.strip()}"
                    )
                # LibreOffice saves the file as <stem>.pdf in the output dir.
                generated = pdf_path.parent / (docx_path.stem + ".pdf")
                if generated != pdf_path and generated.exists():
                    generated.rename(pdf_path)
            else:
                raise NotImplementedError(
                    "Only LibreOffice-based PDF export is supported. "
                    "Set use_libreoffice=True."
                )

            return {
                "path": p.path,
                "pdf_path": str(pdf_path),
                "exported": True,
            }

        return await asyncio.to_thread(_export)

    # ------------------------------------------------------------------
    # Manifest
    # ------------------------------------------------------------------

    def get_manifest(self) -> ModuleManifest:
        return ModuleManifest(
            module_id=self.MODULE_ID,
            version=self.VERSION,
            description=(
                "Create, read, edit, and export Microsoft Word (.docx) documents. "
                "Supports paragraphs, tables, images, headers/footers, hyperlinks, "
                "bookmarks, table of contents, find-and-replace, and PDF export."
            ),
            tags=["word", "docx", "document", "office"],
            dependencies=["python-docx>=1.0"],
            declared_permissions=["filesystem_read", "filesystem_write"],
            actions=[
                ActionSpec(
                    name="open_document",
                    description="Open a .docx file and return its structure summary.",
                    params=[
                        ParamSpec("path", "string", "Path to the .docx file."),
                    ],
                    returns="object",
                    returns_description=(
                        '{"path": str, "paragraph_count": int, '
                        '"table_count": int, "section_count": int}'
                    ),
                    tags=["lifecycle"],
                ),
                ActionSpec(
                    name="create_document",
                    description="Create a new blank Word document, set metadata and default font, and save it.",
                    params=[
                        ParamSpec("output_path", "string", "Destination .docx file path."),
                        ParamSpec("title", "string", "Document title.", required=False),
                        ParamSpec("author", "string", "Document author.", required=False),
                        ParamSpec("default_font", "string", "Default font name.", required=False),
                        ParamSpec("default_font_size", "integer", "Default font size in pt.", required=False),
                    ],
                    returns="object",
                    returns_description='{"path": str, "created": bool}',
                    tags=["lifecycle"],
                ),
                ActionSpec(
                    name="save_document",
                    description="Save a cached document to disk.",
                    params=[
                        ParamSpec("path", "string", "Source .docx path (must be open/cached)."),
                        ParamSpec("output_path", "string", "Save-as path.", required=False),
                    ],
                    returns="object",
                    returns_description='{"path": str, "saved": bool}',
                    tags=["lifecycle"],
                ),
                ActionSpec(
                    name="set_document_properties",
                    description="Update core document properties (title, author, subject, etc.).",
                    params=[
                        ParamSpec("path", "string", "Path to the .docx file."),
                        ParamSpec("title", "string", "Document title.", required=False),
                        ParamSpec("author", "string", "Author name.", required=False),
                        ParamSpec("subject", "string", "Document subject.", required=False),
                        ParamSpec("description", "string", "Document description.", required=False),
                        ParamSpec("keywords", "string", "Keywords string.", required=False),
                        ParamSpec("category", "string", "Document category.", required=False),
                    ],
                    tags=["lifecycle"],
                ),
                ActionSpec(
                    name="get_document_meta",
                    description="Return all core_properties of the document.",
                    params=[
                        ParamSpec("path", "string", "Path to the .docx file."),
                    ],
                    returns="object",
                    tags=["lifecycle", "read"],
                ),
                ActionSpec(
                    name="set_margins",
                    description="Set page margins for a section.",
                    params=[
                        ParamSpec("path", "string", "Path to the .docx file."),
                        ParamSpec("top", "number", "Top margin in cm.", required=False, default=2.54),
                        ParamSpec("bottom", "number", "Bottom margin in cm.", required=False, default=2.54),
                        ParamSpec("left", "number", "Left margin in cm.", required=False, default=3.17),
                        ParamSpec("right", "number", "Right margin in cm.", required=False, default=3.17),
                        ParamSpec("section", "integer", "0-indexed section.", required=False, default=0),
                    ],
                    tags=["lifecycle"],
                ),
                ActionSpec(
                    name="set_default_font",
                    description="Change the Normal style font name and/or size for the whole document.",
                    params=[
                        ParamSpec("path", "string", "Path to the .docx file."),
                        ParamSpec("font_name", "string", "Font name, e.g. 'Calibri'."),
                        ParamSpec("font_size", "integer", "Font size in pt.", required=False),
                    ],
                    tags=["lifecycle"],
                ),
                ActionSpec(
                    name="read_document",
                    description=(
                        "Extract full document content — paragraphs, optionally tables "
                        "and headers/footers."
                    ),
                    params=[
                        ParamSpec("path", "string", "Path to the .docx file."),
                        ParamSpec("include_tables", "boolean", "Include table data.", required=False, default=True),
                        ParamSpec("include_headers_footers", "boolean", "Include headers/footers.", required=False, default=False),
                    ],
                    returns="object",
                    tags=["read"],
                ),
                ActionSpec(
                    name="list_paragraphs",
                    description="List all paragraphs with their index, text, style and empty status.",
                    params=[
                        ParamSpec("path", "string", "Path to the .docx file."),
                        ParamSpec("include_empty", "boolean", "Include empty paragraphs.", required=False, default=False),
                        ParamSpec("style_filter", "string", "Filter by style name.", required=False),
                    ],
                    tags=["read"],
                ),
                ActionSpec(
                    name="list_tables",
                    description="List all tables with index, row/col counts, and first cell preview.",
                    params=[
                        ParamSpec("path", "string", "Path to the .docx file."),
                    ],
                    tags=["read"],
                ),
                ActionSpec(
                    name="extract_text",
                    description="Extract all text from the document as a single string.",
                    params=[
                        ParamSpec("path", "string", "Path to the .docx file."),
                        ParamSpec("separator", "string", "Paragraph separator.", required=False, default="\n"),
                        ParamSpec("include_tables", "boolean", "Include table cell text.", required=False, default=True),
                    ],
                    returns="object",
                    returns_description='{"path": str, "text": str, "length": int}',
                    tags=["read"],
                ),
                ActionSpec(
                    name="count_words",
                    description="Count total words across all paragraph and table cell text.",
                    params=[
                        ParamSpec("path", "string", "Path to the .docx file."),
                    ],
                    returns="object",
                    returns_description='{"path": str, "word_count": int}',
                    tags=["read"],
                ),
                ActionSpec(
                    name="write_paragraph",
                    description=(
                        "Add a paragraph with text, style, and formatting. "
                        "Optionally insert at a specific position."
                    ),
                    params=[
                        ParamSpec("path", "string", "Path to the .docx file."),
                        ParamSpec("text", "string", "Paragraph text."),
                        ParamSpec("style", "string", "Paragraph style name.", required=False, default="Normal"),
                        ParamSpec("bold", "boolean", "Bold text.", required=False, default=False),
                        ParamSpec("italic", "boolean", "Italic text.", required=False, default=False),
                        ParamSpec("underline", "boolean", "Underlined text.", required=False, default=False),
                        ParamSpec("font_name", "string", "Font name.", required=False),
                        ParamSpec("font_size", "integer", "Font size in pt.", required=False),
                        ParamSpec("font_color", "string", "Hex colour e.g. 'FF0000'.", required=False),
                        ParamSpec(
                            "alignment",
                            "string",
                            "Alignment: left, center, right, justify.",
                            required=False,
                            default="left",
                            enum=["left", "center", "right", "justify"],
                        ),
                        ParamSpec("insert_after_paragraph", "integer", "Insert after this paragraph index.", required=False),
                    ],
                    tags=["paragraphs"],
                ),
                ActionSpec(
                    name="format_text",
                    description="Apply character-level formatting to a specific run within a paragraph.",
                    params=[
                        ParamSpec("path", "string", "Path to the .docx file."),
                        ParamSpec("paragraph_index", "integer", "0-indexed paragraph."),
                        ParamSpec("run_index", "integer", "0-indexed run.", required=False, default=0),
                        ParamSpec("bold", "boolean", "Bold.", required=False),
                        ParamSpec("italic", "boolean", "Italic.", required=False),
                        ParamSpec("font_name", "string", "Font name.", required=False),
                        ParamSpec("font_size", "integer", "Font size in pt.", required=False),
                        ParamSpec("font_color", "string", "Hex colour.", required=False),
                    ],
                    tags=["paragraphs"],
                ),
                ActionSpec(
                    name="apply_style",
                    description="Apply a named Word style to a paragraph.",
                    params=[
                        ParamSpec("path", "string", "Path to the .docx file."),
                        ParamSpec("paragraph_index", "integer", "0-indexed paragraph."),
                        ParamSpec("style", "string", "Style name to apply."),
                    ],
                    tags=["paragraphs"],
                ),
                ActionSpec(
                    name="delete_paragraph",
                    description="Remove a paragraph from the document.",
                    params=[
                        ParamSpec("path", "string", "Path to the .docx file."),
                        ParamSpec("paragraph_index", "integer", "0-indexed paragraph to delete."),
                    ],
                    tags=["paragraphs"],
                ),
                ActionSpec(
                    name="insert_page_break",
                    description="Insert a page break, optionally after a specific paragraph.",
                    params=[
                        ParamSpec("path", "string", "Path to the .docx file."),
                        ParamSpec("after_paragraph", "integer", "Insert after this paragraph index.", required=False),
                    ],
                    tags=["paragraphs"],
                ),
                ActionSpec(
                    name="insert_section_break",
                    description="Insert a section break of the given type.",
                    params=[
                        ParamSpec("path", "string", "Path to the .docx file."),
                        ParamSpec(
                            "break_type",
                            "string",
                            "Section break type.",
                            required=False,
                            default="nextPage",
                            enum=["nextPage", "continuous", "evenPage", "oddPage"],
                        ),
                        ParamSpec("after_paragraph", "integer", "Insert after this paragraph index.", required=False),
                    ],
                    tags=["paragraphs"],
                ),
                ActionSpec(
                    name="insert_list",
                    description="Insert a bulleted or numbered list.",
                    params=[
                        ParamSpec("path", "string", "Path to the .docx file."),
                        ParamSpec("items", "array", "List of text items."),
                        ParamSpec(
                            "list_style",
                            "string",
                            "List style: bullet, number, alpha, roman.",
                            required=False,
                            default="bullet",
                            enum=["bullet", "number", "alpha", "roman"],
                        ),
                        ParamSpec("indent_level", "integer", "Indent level (0-8).", required=False, default=0),
                        ParamSpec("insert_after_paragraph", "integer", "Insert after this paragraph index.", required=False),
                    ],
                    tags=["paragraphs"],
                ),
                ActionSpec(
                    name="insert_table",
                    description="Insert a table with given dimensions, optional data, style, and header row.",
                    params=[
                        ParamSpec("path", "string", "Path to the .docx file."),
                        ParamSpec("rows", "integer", "Number of rows."),
                        ParamSpec("cols", "integer", "Number of columns."),
                        ParamSpec("data", "array", "Row-major list of cell values.", required=False),
                        ParamSpec("style", "string", "Table style name.", required=False, default="Table Grid"),
                        ParamSpec("has_header", "boolean", "Mark first row as header.", required=False, default=True),
                        ParamSpec("insert_after_paragraph", "integer", "Insert after this paragraph index.", required=False),
                    ],
                    tags=["tables"],
                ),
                ActionSpec(
                    name="modify_table_cell",
                    description="Update text and formatting of a specific table cell.",
                    params=[
                        ParamSpec("path", "string", "Path to the .docx file."),
                        ParamSpec("table_index", "integer", "0-indexed table."),
                        ParamSpec("row", "integer", "0-indexed row."),
                        ParamSpec("col", "integer", "0-indexed column."),
                        ParamSpec("text", "string", "New cell text.", required=False),
                        ParamSpec("bold", "boolean", "Bold text.", required=False),
                        ParamSpec("font_size", "integer", "Font size in pt.", required=False),
                        ParamSpec("bg_color", "string", "Background hex colour.", required=False),
                    ],
                    tags=["tables"],
                ),
                ActionSpec(
                    name="add_table_row",
                    description="Append a new row to an existing table.",
                    params=[
                        ParamSpec("path", "string", "Path to the .docx file."),
                        ParamSpec("table_index", "integer", "0-indexed table."),
                        ParamSpec("data", "array", "Cell values for the new row."),
                    ],
                    tags=["tables"],
                ),
                ActionSpec(
                    name="insert_image",
                    description="Insert an image into the document with optional width, caption, and alignment.",
                    params=[
                        ParamSpec("path", "string", "Path to the .docx file."),
                        ParamSpec("image_path", "string", "Path to the image file."),
                        ParamSpec("width_cm", "number", "Image width in cm.", required=False),
                        ParamSpec("height_cm", "number", "Image height in cm.", required=False),
                        ParamSpec("caption", "string", "Optional caption text.", required=False),
                        ParamSpec("alignment", "string", "Alignment: left, center, right.", required=False, default="left"),
                        ParamSpec("insert_after_paragraph", "integer", "Insert after this paragraph index.", required=False),
                    ],
                    tags=["rich_content"],
                ),
                ActionSpec(
                    name="insert_hyperlink",
                    description="Add a clickable hyperlink to an existing paragraph.",
                    params=[
                        ParamSpec("path", "string", "Path to the .docx file."),
                        ParamSpec("paragraph_index", "integer", "0-indexed target paragraph."),
                        ParamSpec("text", "string", "Display text for the link."),
                        ParamSpec("url", "string", "URL for the hyperlink."),
                        ParamSpec("font_name", "string", "Font name.", required=False),
                        ParamSpec("font_size", "integer", "Font size in pt.", required=False),
                    ],
                    tags=["rich_content"],
                ),
                ActionSpec(
                    name="add_bookmark",
                    description="Add a named bookmark to a paragraph.",
                    params=[
                        ParamSpec("path", "string", "Path to the .docx file."),
                        ParamSpec("paragraph_index", "integer", "0-indexed paragraph."),
                        ParamSpec("name", "string", "Bookmark name (letters, digits, underscore)."),
                    ],
                    tags=["rich_content"],
                ),
                ActionSpec(
                    name="add_comment",
                    description=(
                        "Add a comment to a paragraph. Implemented as a visible note paragraph "
                        "since python-docx does not support native comment XML."
                    ),
                    params=[
                        ParamSpec("path", "string", "Path to the .docx file."),
                        ParamSpec("paragraph_index", "integer", "0-indexed paragraph."),
                        ParamSpec("text", "string", "Comment text."),
                        ParamSpec("author", "string", "Comment author.", required=False, default="LLMOS Bridge"),
                    ],
                    tags=["rich_content"],
                ),
                ActionSpec(
                    name="insert_toc",
                    description=(
                        "Insert a Table of Contents field code. "
                        "Requires document to be opened in Word and TOC refreshed (F9)."
                    ),
                    params=[
                        ParamSpec("path", "string", "Path to the .docx file."),
                        ParamSpec("title", "string", "TOC heading text.", required=False, default="Contents"),
                        ParamSpec("max_depth", "integer", "Maximum heading depth (1-9).", required=False, default=3),
                        ParamSpec("insert_after_paragraph", "integer", "Insert after this paragraph index.", required=False),
                    ],
                    tags=["rich_content"],
                ),
                ActionSpec(
                    name="add_header_footer",
                    description="Set header and/or footer text for a section, with optional page numbers.",
                    params=[
                        ParamSpec("path", "string", "Path to the .docx file."),
                        ParamSpec("header_text", "string", "Header text.", required=False),
                        ParamSpec("footer_text", "string", "Footer text.", required=False),
                        ParamSpec("section", "integer", "0-indexed section.", required=False, default=0),
                        ParamSpec("page_numbers", "boolean", "Add page numbers to footer.", required=False, default=False),
                        ParamSpec("alignment", "string", "Text alignment: left, center, right.", required=False, default="center"),
                    ],
                    tags=["header_footer"],
                ),
                ActionSpec(
                    name="find_replace",
                    description="Find and replace all occurrences of a text string throughout the document.",
                    params=[
                        ParamSpec("path", "string", "Path to the .docx file."),
                        ParamSpec("find", "string", "Text to search for."),
                        ParamSpec("replace", "string", "Replacement text."),
                        ParamSpec("case_sensitive", "boolean", "Case-sensitive search.", required=False, default=False),
                        ParamSpec("whole_word", "boolean", "Match whole words only.", required=False, default=False),
                    ],
                    returns="object",
                    returns_description='{"path": str, "find": str, "replace": str, "replacements_made": int}',
                    tags=["search"],
                ),
                ActionSpec(
                    name="export_to_pdf",
                    description="Export the document to PDF using LibreOffice headless.",
                    params=[
                        ParamSpec("path", "string", "Path to the .docx file."),
                        ParamSpec("output_path", "string", "Destination PDF path."),
                        ParamSpec("use_libreoffice", "boolean", "Use LibreOffice for conversion.", required=False, default=True),
                    ],
                    returns="object",
                    returns_description='{"path": str, "pdf_path": str, "exported": bool}',
                    permission_required="local_worker",
                    tags=["export"],
                ),
            ],
        )
